"""Tests for docx_writer — render sample Markdown into the branded template and
reopen with python-docx to assert the cover and native structure round-trip."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from docx_writer import write_docx


_SECTION = "\n".join([
    "## Summary",
    "",
    "An assessment with **bold** text and an `inline code` span.",
    "",
    "| Criticality | Number of instances |",
    "|---|---|",
    "| Red Flag | 1 |",
    "| Suggestion | 0 |",
    "",
    "## Detailed Assessment",
    "",
    "### SQL Server",
    "",
    "**1. Are the connections secured? — Red Flag**",
    "Answer: No",
    "",
    "## Outstanding Questions",
    "",
    "- Network latency was not provided.",
])


def _write(tmp_path, **kw):
    out = tmp_path / "report.docx"
    write_docx([_SECTION], out, organisation="Acme Corporation",
               assessment_date="04 July 2026", **kw)
    return out


def test_write_docx_uses_template_cover(tmp_path: Path) -> None:
    """The branded template is used: the cover carries the Organisation (Title)
    and the fixed Subtitle, and the assessment date."""
    doc = Document(str(_write(tmp_path)))
    by_style = {p.style.name: p.text for p in doc.paragraphs
                if p.style.name in ("Title", "Subtitle", "Subtitle2")}
    assert by_style.get("Title") == "Acme Corporation"
    assert by_style.get("Subtitle") == "Online Technical Infrastructure Assessment Report"
    assert "04 July 2026" in by_style.get("Subtitle2", "")
    # The cover graphic / logo / theme are carried by the template parts.
    assert len(doc.sections) >= 2                      # cover + main body


def test_write_docx_renders_body_with_branded_styles(tmp_path: Path) -> None:
    """`##` → Heading 1, `###` → Heading 2; table + bold runs round-trip."""
    doc = Document(str(_write(tmp_path)))
    h1 = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    h2 = [p.text for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert "Summary" in h1 and "Detailed Assessment" in h1
    assert "SQL Server" in h2
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert t.rows[0].cells[0].text == "Criticality"
    assert t.rows[1].cells[0].text == "Red Flag"
    # A bold run from a **bold** block lead exists in the body.
    assert any(r.bold for p in doc.paragraphs for r in p.runs)


def test_write_docx_falls_back_when_template_missing(tmp_path: Path) -> None:
    """A missing template degrades to a blank document — still writes a usable
    report without raising."""
    out = tmp_path / "plain.docx"
    write_docx([_SECTION], out, organisation="Acme",
               assessment_date="04 July 2026",
               template_path=tmp_path / "does_not_exist.docx")
    assert out.exists()
    doc = Document(str(out))
    assert any("Are the connections secured?" in p.text for p in doc.paragraphs)


def test_write_docx_creates_parent_dir(tmp_path: Path) -> None:
    out = tmp_path / "nested" / "deep" / "r.docx"
    write_docx(["## S\ncontent"], out, organisation="Acme", assessment_date="")
    assert out.exists()
