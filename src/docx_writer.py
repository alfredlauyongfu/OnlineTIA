"""Render a TIA report into a branded Microsoft Word (.docx) document.

This writer opens the SS&C / Blue Prism house-style template
(`assets/tia_template.docx` — a branded cover page, running header with logo,
"Commercial in Confidence" + page-number footer, Arial Nova theme and named
Heading styles), fills the cover with the customer's Organisation and the
assessment date, then appends the report body — mapping each Markdown construct
in the per-section content to a native Word element using the template's named
styles (Heading 1/2, tables, lists, bold/code runs). It is INDEPENDENT of the
written `.md` file.

If the template is missing or unreadable it falls back to a blank document, and
the generator calls `write_docx` inside a best-effort try/except, so any docx
problem degrades to a logged warning without affecting the `.md` output or the
pipeline.

Only the limited, well-formed set of Markdown constructs the TIA reports
actually use is handled; anything unrecognised becomes a plain paragraph so
rendering never raises on unexpected input.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.shared import Pt
from docx.text.paragraph import Paragraph


logger = logging.getLogger(__name__)

# The branded template ships in the repo alongside src/ (repo_root/assets/).
DEFAULT_TEMPLATE = Path(__file__).resolve().parent.parent / "assets" / "tia_template.docx"

# The fixed cover subtitle (baked into the template; here for the blank fallback).
COVER_SUBTITLE = "Online Technical Infrastructure Assessment Report"

# A Markdown heading line: 1-6 leading '#', then the text.
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
# A table separator row, e.g. |---|:--:|---| (dashes, colons, pipes, spaces).
_TABLE_SEP_RE = re.compile(r"^\s*\|?[\s:|-]*-[\s:|-]*\|?\s*$")
# Unordered list item: '-' or '*' bullet.
_ULIST_RE = re.compile(r"^\s*[-*]\s+(.*)$")
# Ordered list item: '1.' / '2)' etc.
_OLIST_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")
# Inline split points: **bold** or `code` (kept as delimiters via capture).
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
_MONO_FONT = "Consolas"

# Preferred style names per list/quote kind, most-branded first. The template
# carries the "(BP)" styles; the trailing names are python-docx defaults for the
# blank-fallback path. `_styled_paragraph` tries each and degrades to plain.
_BULLET_STYLES = ("Bullet List (BP)", "Bullet", "List Bullet")
_NUMBER_STYLES = ("Numbered List (BP)", "List Number")
_QUOTE_STYLES = ("Note", "Quote", "Intense Quote")
_TABLE_STYLES = ("Table Grid",)


def _styled_paragraph(doc: _Document, style_candidates: tuple[str, ...]) -> Paragraph:
    """Add a paragraph and apply the first available style from
    `style_candidates`; if none exist in the document, leave it Normal. Applying
    a missing style raises, so this keeps rendering robust across templates."""
    p = doc.add_paragraph()
    for style in style_candidates:
        try:
            p.style = style
            break
        except (KeyError, ValueError):
            continue
    return p


def write_docx(
    section_texts: list[str],
    docx_path: Path,
    *,
    organisation: str,
    assessment_date: str,
    template_path: Path | None = DEFAULT_TEMPLATE,
) -> None:
    """Build a branded .docx at `docx_path` from the report's section bodies.

    `section_texts` are the per-section Markdown strings (already stripped of
    code fences) in document order. The cover shows `organisation` (Title) and
    `assessment_date` (below the fixed subtitle). If the template can't be
    opened, a blank document is used instead (still writes a usable report).
    """
    doc, templated = _open_template(template_path)
    if templated:
        _fill_cover(doc, organisation, assessment_date)
    else:
        # Blank fallback — no cover art, but still a titled, readable report.
        doc.add_heading(organisation or "Customer", level=0)
        doc.add_paragraph(COVER_SUBTITLE)
        if assessment_date:
            doc.add_paragraph(assessment_date)

    for section in section_texts:
        _render_block(doc, section)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


# ---- template + cover ----

def _open_template(template_path: Path | None) -> tuple[_Document, bool]:
    """Open the branded template; return (doc, templated). Falls back to a
    blank `Document()` (with a warning) if the template is absent/unreadable."""
    try:
        if template_path is not None and Path(template_path).is_file():
            return Document(str(template_path)), True
        logger.warning(
            "TIA docx template not found at %s; using a blank document", template_path
        )
    except Exception as exc:  # noqa: BLE001 - any load error degrades gracefully
        logger.warning("TIA docx template unreadable (%s); using a blank document", exc)
    return Document(), False


def _fill_cover(doc: _Document, organisation: str, assessment_date: str) -> None:
    """Set the cover Title (Organisation) and Subtitle2 (date) placeholders; the
    Subtitle is fixed in the template. No-op for a style that isn't present."""
    _set_cover_text(doc, "Title", organisation or "Customer")
    _set_cover_text(
        doc, "Subtitle2",
        f"Assessment date: {assessment_date}" if assessment_date else "",
    )


def _set_cover_text(doc: _Document, style_name: str, text: str) -> None:
    """Replace the text of the first paragraph styled `style_name`, preserving
    that paragraph's style (so the branded formatting is kept)."""
    for p in doc.paragraphs:
        if p.style is not None and p.style.name == style_name:
            for r in list(p.runs):
                r.text = ""
            if p.runs:
                p.runs[0].text = text
            else:
                p.add_run(text)
            return


# ---- internals ----

def _render_block(doc: _Document, markdown: str) -> None:
    """Render one Markdown block (title or section) into `doc`."""
    lines = markdown.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line — nothing to add (Word paragraphs already separate).
        if not stripped:
            i += 1
            continue

        # Thematic break: a line of only dashes (not a table separator, which
        # is handled in the table branch). Render as a thin spacer paragraph.
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # Heading. The report's top sections use `##` and category subsections
        # use `###`; map `##` → Word Heading 1, `###` → Heading 2, so the branded
        # Heading 1/2 styles apply (the document title lives on the cover, not
        # in the body).
        m = _HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2).strip(), level=max(1, min(level - 1, 4)))
            i += 1
            continue

        # Table: a '|' line immediately followed by a separator row.
        if "|" in line and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1]):
            i = _render_table(doc, lines, i)
            continue

        # Blockquote.
        if stripped.startswith(">"):
            p = _styled_paragraph(doc, _QUOTE_STYLES)
            _add_runs(p, stripped.lstrip(">").strip())
            i += 1
            continue

        # Unordered list.
        m = _ULIST_RE.match(line)
        if m:
            p = _styled_paragraph(doc, _BULLET_STYLES)
            _add_runs(p, m.group(1).strip())
            i += 1
            continue

        # Ordered list.
        m = _OLIST_RE.match(line)
        if m:
            p = _styled_paragraph(doc, _NUMBER_STYLES)
            _add_runs(p, m.group(1).strip())
            i += 1
            continue

        # Plain paragraph.
        p = doc.add_paragraph()
        _add_runs(p, stripped)
        i += 1


def _render_table(doc: _Document, lines: list[str], start: int) -> int:
    """Render a GFM pipe table starting at `lines[start]` (the header row,
    with `lines[start+1]` the separator). Returns the index just past the
    table."""
    header = _split_row(lines[start])
    body: list[list[str]] = []
    i = start + 2  # skip header + separator
    while i < len(lines) and "|" in lines[i] and lines[i].strip():
        body.append(_split_row(lines[i]))
        i += 1

    ncols = max([len(header)] + [len(r) for r in body]) or 1
    table = doc.add_table(rows=1, cols=ncols)
    for style in _TABLE_STYLES:
        try:
            table.style = style
            break
        except (KeyError, ValueError):
            continue  # style not in this template — try the next / leave unstyled

    # Header row (bold).
    hdr_cells = table.rows[0].cells
    for c in range(ncols):
        text = header[c] if c < len(header) else ""
        para = hdr_cells[c].paragraphs[0]
        _add_runs(para, text, bold=True)

    # Body rows.
    for row in body:
        cells = table.add_row().cells
        for c in range(ncols):
            text = row[c] if c < len(row) else ""
            _add_runs(cells[c].paragraphs[0], text)

    return i


def _split_row(line: str) -> list[str]:
    """Split a Markdown table row into trimmed cell strings, dropping the
    leading/trailing pipe-induced empties."""
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _add_runs(paragraph: Paragraph, text: str, bold: bool = False) -> None:
    """Append `text` to `paragraph`, honouring inline **bold** and `code`.

    `bold=True` makes the whole segment bold (used for table headers); inline
    `**...**` still forces bold and `` `...` `` renders monospaced.
    """
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = _MONO_FONT
            run.font.size = Pt(10)
            if bold:
                run.bold = True
        else:
            run = paragraph.add_run(part)
            if bold:
                run.bold = True
